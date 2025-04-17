from typing import List, Tuple
from dataclasses import dataclass
import re
import docx
from docx.text.paragraph import Paragraph

@dataclass
class QuestionBlock:
    question: str
    answers: List[str]
    correct_answer: int

def clean_answer(text: str) -> str:
    cleaned_text = text.strip()
    cleaned_text = re.sub(r"^(А|Б|В|Г|а|б|в|г|\-)\)", "", cleaned_text).strip()
    cleaned_text = re.sub(r"[.;]$", "", cleaned_text).strip()
    return cleaned_text

def get_question_block_by_full_text(paragraph: Paragraph) -> QuestionBlock:
    current_line = ""
    bold_found = False
    is_first_line = True
    question = ""
    answers = []
    correct_answer = -1

    for run in paragraph.runs:
        parts = run.text.split("\n")
        for i, part in enumerate(parts):
            current_line += part
            if run.bold:
                bold_found = True
            if i < len(parts) - 1:
                if is_first_line:
                    question = current_line
                    is_first_line = False
                else:
                    answers.append(clean_answer(current_line))
                    correct_answer = len(answers) - 1 if bold_found else correct_answer
                current_line = ""
                bold_found = False

    if not is_first_line and current_line.strip():
        answers.append(clean_answer(current_line))
        correct_answer = len(answers) - 1 if bold_found else correct_answer

    return QuestionBlock(
        question=question,
        answers=answers,
        correct_answer=correct_answer
    )

def has_paragraph_numbering(paragraph: Paragraph) -> bool:
    if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
        ilvl = paragraph._p.pPr.numPr.ilvl
        if ilvl is not None:
            return True
    return False

def get_question_block_by_lines(paragraphs: List[Paragraph], start_index: int) -> Tuple[QuestionBlock, int]:    
    question = ""
    answers: List[str] = []
    correct_answer = -1
    index = start_index

    # Зчитати питання
    para = paragraphs[index]

    if has_paragraph_numbering(para):
        text = para.text.strip()
        start_ilvl = para._p.pPr.numPr.ilvl.val
        question = text

        index += 1

        # Зчитати відповіді
        while index < len(paragraphs):
            para = paragraphs[index]
            if has_paragraph_numbering(para) is False:
                break
            
            text = para.text.strip()
            ilvl = para._p.pPr.numPr.ilvl.val

            if ilvl <= start_ilvl:
                break

            is_bold = any(run.bold for run in para.runs if run.text.strip())
            cleaned = clean_answer(text)
            answers.append(cleaned)

            if is_bold:
                correct_answer = len(answers) - 1

            index += 1
    else:
        index += 1

    return QuestionBlock(question=question, answers=answers, correct_answer=correct_answer), index

def get_question_blocks(doc: docx.Document) -> List[QuestionBlock]:
    question_blocks: List[QuestionBlock] = []
    paragraph_id = 0
    while paragraph_id < len(doc.paragraphs):
        block = get_question_block_by_full_text(doc.paragraphs[paragraph_id])
        if not block.answers:
            block, paragraph_id = get_question_block_by_lines(doc.paragraphs, paragraph_id)
        else:
            paragraph_id += 1

        if block.answers:
            question_blocks.append(block)
    return question_blocks
